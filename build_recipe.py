import os
import glob
import subprocess
import json
from docx import Document
import openai

# --- НАСТРОЙКИ ПУТЕЙ И API ---
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
INPUT_DIR = os.path.join(BASE_DIR, 'input')
OUTPUT_DIR = os.path.join(BASE_DIR, 'output')
TEMPLATES_DIR = os.path.join(BASE_DIR, 'templates')

TEMPLATE_NAME = 'recipe_template.indd'
TEMPLATE_PATH = os.path.join(TEMPLATES_DIR, TEMPLATE_NAME)

# API Ключ для работы с AI
OPENAI_API_KEY = os.environ.get("OPENAI_API_KEY", "")

if not OPENAI_API_KEY:
    print("[ERROR] Не задан OPENAI_API_KEY в переменных окружения.")
    print("Выполните команду в терминале перед запуском: export OPENAI_API_KEY='ваш_ключ'")
    exit(1)

client = openai.OpenAI(api_key=OPENAI_API_KEY)


def check_environment():
    """Проверяет наличие необходимых папок и шаблона."""
    for directory in [INPUT_DIR, OUTPUT_DIR, TEMPLATES_DIR]:
        if not os.path.exists(directory):
            os.makedirs(directory)
            print(f"[INFO] Создана папка: {directory}")

    if not os.path.exists(TEMPLATE_PATH):
        print(f"[ERROR] Шаблон не найден: {TEMPLATE_PATH}")
        print("Пожалуйста, поместите InDesign шаблон в папку templates/.")
        return False
    return True

def extract_first_image_from_docx(docx_path, output_dir, base_name):
    """Извлекает первую найденную картинку из .docx файла."""
    try:
        doc = Document(docx_path)
        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:
                img_blob = rel.target_part.blob
                ext = rel.target_ref.split('.')[-1]
                if ext.lower() not in ['png', 'jpg', 'jpeg']:
                    ext = 'jpg'
                img_path = os.path.join(output_dir, f"{base_name}_extracted.{ext}")
                with open(img_path, "wb") as f:
                    f.write(img_blob)
                return img_path
    except Exception as e:
        print(f"[WARNING] Не удалось извлечь картинку из {docx_path}: {e}")
    return None

def parse_docx_raw(docx_path):
    """Извлекает весь текст из документа без разбивки, просто сырой строкой."""
    try:
        doc = Document(docx_path)
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        return "\n".join(paragraphs)
    except Exception as e:
        print(f"[ERROR] Ошибка чтения {docx_path}: {e}")
        return None

def extract_template_frames_from_indesign():
    """Открывает шаблон InDesign и вытаскивает все ID текстовых фреймов и их текущий текст."""
    print("[INFO] Читаю структуру текстовых фреймов из шаблона InDesign (Идет открытие шаблона, ждите)...")
    jsx_path = os.path.join(BASE_DIR, "temp_extract.jsx")
    json_path = os.path.join(BASE_DIR, "frames.json")
    
    template_js = TEMPLATE_PATH.replace("\\", "/")
    json_js = json_path.replace("\\", "/")
    
    jsx_code = f"""
app.scriptPreferences.userInteractionLevel = UserInteractionLevels.NEVER_INTERACT;
try {{
    var templateFile = new File("{template_js}");
    var doc = app.open(templateFile, false);
    
    var jsonStrs = [];
    for (var i = 0; i < doc.textFrames.length; i++) {{
        var rawText = doc.textFrames[i].contents;
        var cleanText = rawText.replace(/\\\\/g, '\\\\\\\\').replace(/"/g, '\\\\"').replace(/\\n/g, '\\\\n').replace(/\\r/g, '\\\\n').replace(/\\t/g, ' ');
        // Берем только первые 400 символов для контекста, чтобы не перегружать AI
        cleanText = cleanText.substring(0, 400); 
        jsonStrs.push('{{"id":"' + i + '", "text":"' + cleanText + '"}}');
    }}
    
    var outFile = new File("{json_js}");
    outFile.encoding = "UTF-8";
    outFile.open("w");
    outFile.write("[" + jsonStrs.join(",") + "]");
    outFile.close();
    
    doc.close(SaveOptions.NO);
}} catch(e) {{
    var errFile = new File("{BASE_DIR}/indesign_error.txt".replace("\\\\", "/"));
    errFile.encoding = "UTF-8";
    errFile.open("w");
    errFile.write(e.toString() + "\\nLine: " + e.line);
    errFile.close();
}} finally {{
    app.scriptPreferences.userInteractionLevel = UserInteractionLevels.INTERACT_WITH_ALL;
}}
"""
    with open(jsx_path, 'w', encoding='utf-8') as f:
        f.write(jsx_code)
        
    subprocess.run(['osascript', '-e', f'tell application "Adobe InDesign 2026" to do script POSIX file "{jsx_path}" language javascript'], capture_output=True)
    os.remove(jsx_path)
    
    if os.path.exists(json_path):
        with open(json_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
        os.remove(json_path)
        
        # Фильтруем пустые фреймы
        data = [item for item in data if item['text'].strip()]
        return data
    else:
        print("[ERROR] Не удалось получить структуру фреймов. Возможно InDesign закрыт или выдал ошибку.")
        return None

def parse_with_ai(frames_data, raw_recipe_text, few_shot_data):
    """Интеграция с OpenAI для маппинга текста нового рецепта во фреймы шаблона (Few-Shot)."""
    print("[INFO] Отправляю задачу нейросети OpenAI (Идет интеллектуальный парсинг с Базой Знаний)...")
    
    frames_context = json.dumps(frames_data, ensure_ascii=False, indent=2)
    few_shot_context = ""
    
    if few_shot_data:
        # Берем до 5 примеров, чтобы не переполнять токенный лимит
        examples = few_shot_data[:5]
        few_shot_context = "ВОТ ПРИМЕРЫ ИЗ ТВОЕЙ БАЗЫ ЗНАНИЙ (Как профессиональный верстальщик раскладывал тексты раньше):\n\n"
        for i, example in enumerate(examples):
            docx_sample = example.get("source_docx", "")[:800] # Берем срез, чтобы не лопнул промпт
            indd_sample = json.dumps(example.get("target_frames", {}), ensure_ascii=False)
            few_shot_context += f"--- ПРИМЕР {i+1} ---\nСЫРОЙ ТЕКСТ:\n{docx_sample}...\n\nРЕЗУЛЬТАТ (JSON ФРЕЙМОВ):\n{indd_sample}\n\n"
            
    prompt = f"""
Ты — профессиональный верстальщик-редактор. Тебе дают шаблон дизайна, где каждый текстовый фрейм имеет свой строковый ID и свой текущий текст (текст прошлого рецепта). 
Также тебе дают текст нового рецепта.

Твоя задача — расставить части нового рецепта по фреймам шаблона с 100% хирургической точностью.
Никакой отсебятины. Копируй текст нового рецепта символ в символ.

ВАЖНЫЕ ПРАВИЛА ВЕРСТКИ ДЛЯ INDESIGN:
1. Ингредиенты [СУПЕР КРИТИЧНО]: 
   - В шаблоне сейчас забиты старые ингредиенты: "Филе куриное", "Бекон", "Шпажки", "Томаты протертые", "Птитим" и т.д.
   - ТВОЯ КРИТИЧЕСКАЯ ЗАДАЧА — ПОЛНОСТЬЮ УНИЧТОЖИТЬ ИХ. Ты НЕ ИМЕЕШЬ ПРАВА возвращать строку, содержащую "Филе куриное" или "Бекон", если их нет в НОВОМ рецепте.
   - Раздели список ингредиентов пополам и заполни обе колонки. Если влезли в одну — для второй верни пустую строку "".
   - [ВЫРАВНИВАНИЕ]: Ингредиенты должны быть выровнены по правому краю с точками! НО НЕ пиши точки сам! Обязательно вставляй СИМВОЛ ТАБУЛЯЦИИ `\\t` между названием ингредиента и весом. Пример: `Картофель мини\\t400 г`. InDesign сам растянет точки до конца строки благодаря табуляции. 
   - [МЕЖСТРОЧНЫЕ ОТСТУПЫ]: КАТЕГОРИЧЕСКИ запрещено делать переносы строки (`\\n`) внутри одного параграфа/ингредиента, даже если он длинный. Напиши длинный ингредиент в одну строку: `Сливочное масло с беконом\\t40 г`. Несоблюдение этого сломает интервалы!
2. КБЖУ:
   - В блоке КБЖУ (Калорийность) между словом и цифрой обязательно ставь точки, чтобы выровнять значения. Пример: `Белки..........................7,5\\nЖиры.............................5.0`
3. Шаги приготовления и нумерация:
   - ОДИН перенос строки `\\n` = ОДИН новый номер шага!
   - КАТЕГОРИЧЕСКИ ЗАПРЕЩАЕТСЯ делать пустые строки (`\\n\\n`) между шагами.
   - КАТЕГОРИЧЕСКИ ЗАПРЕЩАЕТСЯ делать перенос строки ВНУТРИ одного шага.
   - Не пиши сами цифры шагов ("1.", "2.") вручную, пиши только сам текст.
   - У каждого шага в конце ставится ОДИН перенос строки `\\n`.
4. Номер сборки заказа: Обязательно найди в DOCX одиночную цифру сборки (например, "7") и подставь её в соответствующий фрейм.
5. ЗАПРЕТЫ:
   - КАТЕГОРИЧЕСКИ ЗАПРЕЩАЕТСЯ писать слово "РЕЦЕПТ" или "Шаги" перед началом шагов.
   - ДЛЯ СПИСКА ПОСУДЫ/ИНВЕНТАРЯ: Копируй весь инвентарь (кастрюли, миски, сковорода и т.д.) из DOCX абсолютно ЦЕЛИКОМ, не пропуская ни единого слова! Если список длинный — копируй его полностью, как есть в DOCX. Не отрезай слова. Оставляй общие статичные фреймы (соцсети) без изменений.

{few_shot_context}

--- ТЕПЕРЬ ТВОЯ ОЧЕРЕДЬ ---
Структура шаблона (текущие фреймы JSON):
{frames_context}

Текст НОВОГО рецепта:
{raw_recipe_text}

ВЕРНИ СТРОГО JSON-СЛОВАРЬ (ключ: ID фрейма, значение: новый текст). 
Если фрейм является СТАТИЧНОЙ НАДПИСЬЮ и его не надо менять, верни старый текст (копию). 
ВНИМАНИЕ 1: Вторая колонка ингредиентов НЕ ЯВЛЯЕТСЯ статичной! Если все продукты влезли в одну колонку, ТЫ ОБЯЗАН обнулить вторую (вернуть `""`).
ВНИМАНИЕ 2: НЕ ПУТАЙ ИНГРЕДИЕНТЫ И ИНВЕНТАРЬ ПРИ ЗАПОЛНЕНИИ КОЛОНОК! Ингредиенты — это продукты (мясо, овощи). А инвентарь (кастрюли, сковороды) пиши СТРОГО в самый нижний фрейм (где сейчас "Кастрюля, дуршлаг..."). Никогда не пиши кастрюли в колонку с ингредиентами!

Никакого markdown, никакого введения, только чистый JSON-массив объекта, где ключи - строки-числа.
Пример ответа: {{"0": "Срок годности", "1": "Шашлычок", "2": "Томаты.."}}
"""
    try:
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": "You output only minified raw JSON object. Use double quotes. No \u0060\u0060\u0060json block around it."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.1,
            timeout=45
        )
        result_text = response.choices[0].message.content.strip()
        
        # Очистка markdown если есть
        if result_text.startswith("```json"):
            result_text = result_text[7:]
        if result_text.startswith("```"):
            result_text = result_text[3:]
        if result_text.endswith("```"):
            result_text = result_text[:-3]
            
        return json.loads(result_text.strip())
    except Exception as e:
        print(f"[ERROR] Ошибка при парсинге через OpenAI API: {e}")
        return None

def generate_write_jsx(filename_base, mapped_data, image_path, output_indd, output_pdf):
    """Генерирует финальный скрипт, который вставляет новые тексты по ID фреймов."""
    
    template_js = TEMPLATE_PATH.replace("\\", "/")
    image_js = image_path.replace("\\", "/") if image_path else ""
    output_indd_js = output_indd.replace("\\", "/")
    output_pdf_js = output_pdf.replace("\\", "/")

    # Формируем JS-объект для обновления фреймов. Мы должны заменить \n на \r для InDesign
    # И экранировать кавычки
    import re
    updates_js_parts = []
    for frame_id, new_text in mapped_data.items():
        if isinstance(new_text, str):
            clean_text = new_text.strip()
            # Убираем случайное слово РЕЦЕПТ в начале текстового блока (если ИИ сглючил)
            clean_text = re.sub(r'(?i)^(рецепт|шаги)[:\n\s]*', '', clean_text).strip()
            # Заменяем двойные переносы на одинарные, чтобы не было пустых строк с номером
            clean_text = re.sub(r'\n{2,}', '\n', clean_text)
            
            # ЖЕСТКАЯ ЗАМЕНА ТОЧЕК НА ТАБУЛЯЦИЮ (Для красивого центрирования в InDesign)
            # Даже если ИИ накидал точек, Python заменит их на \t
            clean_text = re.sub(r'\s*\.{2,}\s*', '\t', clean_text)
            
            new_text = clean_text + "\n" if clean_text else ""
            
        # InDesign использует \r
        safe_text = str(new_text).replace('"', '\\"').replace("'", "\\'").replace('\n', '\\r').replace('\r\r', '\\r')
        updates_js_parts.append(f'"{frame_id}": "{safe_text}"')
    
    updates_json_str = "{" + ", ".join(updates_js_parts) + "}"

    jsx_code = f"""
app.scriptPreferences.userInteractionLevel = UserInteractionLevels.NEVER_INTERACT;
try {{
    var templateFile = new File("{template_js}");
    var doc = app.open(templateFile, false);
    
    var updates = {updates_json_str};
    
    // 1. Обновление текстов
    for (var key in updates) {{
        var idx = parseInt(key);
        if (idx >= 0 && idx < doc.textFrames.length) {{
            var frame = doc.textFrames[idx];
            if (frame.locked) frame.locked = false;
            frame.contents = updates[key];
        }}
    }}
    
    // 2. Вставка картинки (берем самый большой Rectangle, а остальные большие - удаляем, чтобы не было дублей старых фото)
    var imageFileStr = "{image_js}";
    if (imageFileStr !== "") {{
        var imageFile = new File(imageFileStr);
        if (imageFile.exists) {{
            var photoFrame = null;
            var maxArea = 0;
            var allLargeRects = [];
            
            for (var i = 0; i < doc.rectangles.length; i++) {{
                var rect = doc.rectangles[i];
                var b = rect.geometricBounds;
                var area = (b[2] - b[0]) * (b[3] - b[1]);
                if (area > 10000) {{
                    allLargeRects.push(rect);
                }}
                
                // Приоритет фреймам, внутри которых УЖЕ ЕСТЬ картинка (template placeholder)
                try {{
                    if (rect.graphics.length > 0 && area > 10000) {{
                        photoFrame = rect;
                    }}
                }} catch(e) {{}}
                
                // Фолбэк на просто самый большой, если картинок вообще нет
                if (!photoFrame && area > maxArea) {{
                    maxArea = area;
                    photoFrame = rect;
                }}
            }}
            
            // Удаляем все остальные большие фреймы (подложки со старыми фото), чтобы они не торчали снизу
            for (var k = 0; k < allLargeRects.length; k++) {{
                if (allLargeRects[k] !== photoFrame) {{
                    try {{
                        if (allLargeRects[k].graphics.length > 0) {{
                            if (allLargeRects[k].locked) allLargeRects[k].locked = false;
                            allLargeRects[k].remove();
                        }}
                    }} catch(e) {{}}
                }}
            }}
            
            if (photoFrame) {{
                if (photoFrame.locked) photoFrame.locked = false;
                photoFrame.place(imageFile);
                photoFrame.fit(FitOptions.FILL_PROPORTIONALLY);
                photoFrame.fit(FitOptions.CENTER_CONTENT);
                
                // --- СМАРТ-ОБРЕЗКА ФОТО (ДЛЯ 5мм ОТСТУПА ОТ ЗАГОЛОВКА) ---
                var titleTopY = 1000;
                var maxTitleSize = 0;
                for (var j = 0; j < doc.textFrames.length; j++) {{
                    var tFrame = doc.textFrames[j];
                    var tb = tFrame.geometricBounds;
                    // Ищем фрейм слева (X1 < 50), сверху/посередине (Y1 > 40 && Y1 < 150)
                    if (tb[1] < 50 && tb[0] > 40 && tb[0] < 150) {{
                        var pSize = 0;
                        try {{
                            if (tFrame.paragraphs.length > 0) {{
                                pSize = tFrame.paragraphs[0].pointSize;
                            }}
                        }} catch(e) {{}}
                        
                        // Заголовок - это самый крупный текст в этой зоне (обычно > 18pt)
                        if (pSize > maxTitleSize && pSize >= 15) {{
                            maxTitleSize = pSize;
                            titleTopY = tb[0];
                        }}
                    }}
                }}
                
                if (titleTopY < 1000) {{
                    var newBottom = titleTopY - 5; // Даем 5 мм отступа
                    var pb = photoFrame.geometricBounds;
                    if (newBottom > pb[0] + 30) {{ // Не сжимаем до 0 (оставляем хотя бы 30мм высоты)
                        photoFrame.geometricBounds = [pb[0], pb[1], newBottom, pb[3]];
                        photoFrame.fit(FitOptions.FILL_PROPORTIONALLY);
                        photoFrame.fit(FitOptions.CENTER_CONTENT);
                    }}
                }}
                // ------------------------------------------------
            }}

        }}
    }}
    
    // 3. Сохранение файла
    var outInddFile = new File("{output_indd_js}");
    doc.save(outInddFile);
    
    // 4. Экспорт в PDF
    var outPdfFile = new File("{output_pdf_js}");
    var pdfPreset = app.pdfExportPresets.item("[High Quality Print]");
    if (!pdfPreset.isValid) {{
        pdfPreset = app.pdfExportPresets.firstItem();
    }}
    doc.exportFile(ExportFormat.PDF_TYPE, outPdfFile, false, pdfPreset);
    
    doc.close(SaveOptions.NO);
}} catch (e) {{
    var errFile = new File("{BASE_DIR}/indesign_error.txt".replace("\\\\", "/"));
    errFile.encoding = "UTF-8";
    errFile.open("w");
    errFile.write(e.toString() + "\\nLine: " + e.line);
    errFile.close();
}} finally {{
    app.scriptPreferences.userInteractionLevel = UserInteractionLevels.INTERACT_WITH_ALL;
}}
"""
    return jsx_code

def run_indesign_script(jsx_path):
    """Запускает сгенерированный JSX скрипт через osascript."""
    err_log = os.path.join(BASE_DIR, 'indesign_error.txt')
    if os.path.exists(err_log):
        os.remove(err_log)

    applescript = f'''
    tell application "Adobe InDesign 2026"
        do script POSIX file "{jsx_path}" language javascript
    end tell
    '''
    
    try:
        result = subprocess.run(['osascript', '-e', applescript], capture_output=True)
        if result.returncode != 0:
            err_msg = result.stderr.decode('utf-8', 'replace') if result.stderr else ''
            print(f"[ERROR] Ошибка выполнения AppleScript: {err_msg}")
            return False
            
        if os.path.exists(err_log):
            with open(err_log, 'r', encoding='utf-8', errors='replace') as f:
                error_msg = f.read()
            print(f"[InDesign ERROR] Скрипт InDesign упал:\n{error_msg}")
            return False
            
        return True
    except Exception as e:
        print(f"[ERROR] Сбой запуска osascript: {e}")
        return False


def main():
    if not check_environment():
        return

    # Загружаем базу знаний Few-Shot если она есть
    few_shot_data = None
    dataset_path = os.path.join(BASE_DIR, 'training_dataset.json')
    if os.path.exists(dataset_path):
        try:
            with open(dataset_path, 'r', encoding='utf-8') as f:
                few_shot_data = json.load(f)
            print(f"[INFO] Успешно загружена База Знаний ({len(few_shot_data)} примеров).")
        except Exception as e:
            print(f"[WARNING] Не удалось прочитать базу знаний: {e}")

    # Извлекаем структуру текущего шаблона из InDesign 
    template_frames = extract_template_frames_from_indesign()
    if not template_frames:
        return

    docx_files = glob.glob(os.path.join(INPUT_DIR, '*.docx'))
    if not docx_files:
        print("[INFO] В папке input/ нет .docx файлов для обработки.")
        return

    for docx_path in docx_files:
        filename_base = os.path.splitext(os.path.basename(docx_path))[0]
        print(f"\n[INFO] Обрабатываю рецепт: {filename_base}")
        
        # Получаем сырой текст из ворда
        raw_text = parse_docx_raw(docx_path)
        if not raw_text:
            continue
            
        # Умный маппинг текста во фреймы (МАГИЯ ИИ)
        mapped_data = parse_with_ai(template_frames, raw_text, few_shot_data)
        if not mapped_data:
            print(f"[ERROR] ИИ не смог распарсить рецепт {filename_base}. Пропуск.")
            continue
            
        print("[DEBUG] AI MAPPED DATA:")
        print(json.dumps(mapped_data, indent=2, ensure_ascii=False))
        
        # Добываем картинку (сначала отдельный файл, затем из docx)
        image_path = None
        for ext in ['.jpg', '.jpeg', '.png']:
            possible_path = os.path.join(INPUT_DIR, f"{filename_base}{ext}")
            if os.path.exists(possible_path):
                image_path = possible_path
                break
                
        if not image_path:
            image_path = extract_first_image_from_docx(docx_path, OUTPUT_DIR, filename_base)
            
        # Генерация и запуск финального скрипта
        output_indd = os.path.join(OUTPUT_DIR, f"{filename_base}.indd")
        output_pdf = os.path.join(OUTPUT_DIR, f"{filename_base}.pdf")
        
        jsx_code = generate_write_jsx(filename_base, mapped_data, image_path, output_indd, output_pdf)
        jsx_file_path = os.path.join(BASE_DIR, f"temp_{filename_base}.jsx")
        
        with open(jsx_file_path, 'w', encoding='utf-8') as f:
            f.write(jsx_code)
            
        print(f"[INFO] Запускаю InDesign для вставки данных и экспорта...")
        success = run_indesign_script(jsx_file_path)
        
        if success:
            print(f"[SUCCESS] Готово! Сохранено: output/{filename_base}.pdf")
        
        if os.path.exists(jsx_file_path):
            os.remove(jsx_file_path)

if __name__ == "__main__":
    main()
